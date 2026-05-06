# ============================================================
# DOCX 解析器 — python-docx
# ============================================================

from docx import Document

from app.ai.parsers.base import DocumentParser, ParseResult, count_words
from app.core.logging import logger


class DocxParser(DocumentParser):
    """DOCX 文档解析器"""

    def parse(self, file_path: str) -> ParseResult:
        """解析 DOCX 文档，提取段落和表格文本"""

        doc = Document(file_path)
        paragraphs: list[str] = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        full_text = "\n".join(paragraphs)
        word_count = count_words(full_text)

        # python-docx 没有直接获取页数的方法，按经验估算
        page_count = max(1, word_count // 500)

        logger.info(
            "DOCX parsed",
            extra={
                "file_path": file_path,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "word_count": word_count,
            },
        )

        return ParseResult(
            text=full_text,
            page_count=page_count,
            word_count=word_count,
            is_scanned=False,
        )
